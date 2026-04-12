"""
Resume text extraction service.
Handles PDF (via PyMuPDF) and DOCX (via python-docx) formats.
Returns clean plain text for AI processing.
"""
import io
import logging
from typing import Literal

logger = logging.getLogger(__name__)

# Supported MIME types → their internal format identifier
SUPPORTED_MIME_TYPES: dict[str, str] = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    # Some browsers mis-label DOCX:
    "application/msword": "docx",
    "application/octet-stream": "unknown",  # fallback — sniff from filename
}

MAX_FILE_SIZE_BYTES = 5 * 1024 * 1024  # 5 MB


def _detect_format_from_filename(filename: str) -> Literal["pdf", "docx"] | None:
    """Detect expected format from filename extension."""
    lower = filename.lower()
    if lower.endswith(".pdf"):
        return "pdf"
    if lower.endswith(".docx"):
        return "docx"
    return None


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extract clean plain text from a PDF file.
    Uses PyMuPDF (fitz) for reliable extraction across
    single-column, multi-column, and scanned-like layouts.
    """
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(stream=file_bytes, filetype="pdf")
        pages: list[str] = []

        for page_num in range(len(doc)):
            page = doc[page_num]
            # extract_text("text") gives reading-order plain text
            text = page.get_text("text")
            if text.strip():
                pages.append(text.strip())

        doc.close()
        extracted = "\n\n".join(pages)
        logger.info(f"PDF extraction: {len(pages)} pages, {len(extracted)} chars")
        return extracted

    except ImportError:
        raise RuntimeError(
            "PyMuPDF is not installed. Run: pip install pymupdf"
        )
    except Exception as e:
        logger.error(f"PDF extraction failed: {e}")
        raise RuntimeError(f"Failed to extract text from PDF: {e}")


def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extract clean plain text from a DOCX file.
    Walks all paragraphs and table cells for maximum coverage.
    """
    try:
        from docx import Document

        doc = Document(io.BytesIO(file_bytes))
        parts: list[str] = []

        # 1. Body paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        # 2. Tables (skills grids, education tables, etc.)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    parts.append(row_text)

        extracted = "\n".join(parts)
        logger.info(f"DOCX extraction: {len(parts)} paragraphs/rows, {len(extracted)} chars")
        return extracted

    except ImportError:
        raise RuntimeError(
            "python-docx is not installed. Run: pip install python-docx"
        )
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        raise RuntimeError(f"Failed to extract text from DOCX: {e}")


def extract_text(
    file_bytes: bytes,
    content_type: str,
    filename: str = "",
) -> str:
    """
    Main entry point. Dispatches to the correct extractor.
    Validates file size and format before extraction.

    Args:
        file_bytes: Raw file content in memory.
        content_type: MIME type from the multipart upload header.
        filename: Original filename for extension-based fallback.

    Returns:
        Extracted plain text (UTF-8 string).

    Raises:
        ValueError: For unsupported formats or oversized files.
        RuntimeError: If extraction itself fails.
    """
    # Guard: file size
    if len(file_bytes) > MAX_FILE_SIZE_BYTES:
        raise ValueError(
            f"File too large ({len(file_bytes) // 1024 // 1024}MB). "
            f"Maximum allowed size is 5MB."
        )

    # Guard: empty file
    if len(file_bytes) == 0:
        raise ValueError("Uploaded file is empty.")

    # Resolve format
    fmt = SUPPORTED_MIME_TYPES.get(content_type)
    if fmt == "unknown" or fmt is None:
        # Try filename extension as fallback
        fmt = _detect_format_from_filename(filename)

    if fmt == "pdf":
        return extract_text_from_pdf(file_bytes)
    elif fmt == "docx":
        return extract_text_from_docx(file_bytes)
    else:
        raise ValueError(
            f"Unsupported file format: '{content_type}'. "
            "Only PDF and DOCX files are accepted."
        )
